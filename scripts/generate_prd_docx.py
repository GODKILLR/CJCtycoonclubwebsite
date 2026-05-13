"""Generate PRD.docx from the canonical PRD content.

The markdown version at /PRD.md is the source of truth — this script produces a
.docx mirror for stakeholders who want a Word file. Re-run after editing PRD.md
and update the corresponding block below.

Usage:
    python scripts/generate_prd_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(0xC9, 0xA9, 0x61)  # antique gold
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x68, 0x59)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK
        if level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(12)


def add_title(doc: Document, text: str, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(text)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = INK

    sub = doc.add_paragraph()
    sub_run = sub.add_run(subtitle)
    sub_run.font.size = Pt(11)
    sub_run.italic = True
    sub_run.font.color.rgb = MUTED


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for i, (k, v) in enumerate(rows):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        for c in (cell_k, cell_v):
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_shading(cell_k, "F4F1E8")
        for p in cell_k.paragraphs:
            p.runs and (p.runs[0].font.size or None)
        cell_k.text = ""
        run_k = cell_k.paragraphs[0].add_run(k)
        run_k.bold = True
        run_k.font.size = Pt(10)
        run_k.font.color.rgb = INK
        cell_v.text = ""
        run_v = cell_v.paragraphs[0].add_run(v)
        run_v.font.size = Pt(10)


def add_table(doc: Document, header: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "1A1A1A")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)


def add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BDB69E")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def build() -> Path:
    out = Path(__file__).resolve().parent.parent / "PRD.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Tycoon Club Website", "Product Requirements Document — v1.0, May 2026")

    add_kv_table(doc, [
        ("Project", "Compliance-first rebuild of tycoonclub.io"),
        ("Repository", "github.com/GODKILLR/CJCtycoonclubwebsite"),
        ("Live preview", "cj-ctycoonclubwebsite.vercel.app"),
        ("Status", "Design-review stage — 4 themes shipped, awaiting counsel sign-off and stakeholder selection"),
        ("Document version", "1.0 — May 2026"),
    ])
    add_hr(doc)

    # 1. Background
    add_heading(doc, "1. Background", 1)
    add_paragraph(doc,
        "The previous public-facing Tycoon Club website (tycoonclub.io) carried significant regulatory exposure. "
        "Key claims and structural elements included:")
    for b in [
        "“Get Your FREE Membership” — a deceptive frame for what was effectively a paid-tier referral program.",
        "“$258,026.57 USDT Distributed” jackpot ticker, “Lifetime Earning”, and “sustainable passive income” positioning.",
        "Multi-tier referral structure (“Alliances”, 1-Star/2-Star/.../Legendary Tycoon tiers with 1%, 5%, 10% recurring USDT income based on the activity of recruited members).",
        "Achievement system tied directly to the count of recruited buddies.",
        "No risk disclosure, no jurisdictional gating, no KYC notice, no operating-entity transparency.",
    ]:
        add_bullet(doc, b)
    add_paragraph(doc,
        "In aggregate this framing resembles an unregistered securities offering and a multi-level marketing structure — both of which attract enforcement risk in most major jurisdictions (US: SEC + FTC; EU: ESMA + national consumer protection; UAE: SCA; Singapore: MAS). The SEC’s 2022 enforcement action against the Forsage smart-contract pyramid ($300M+ ponzi/pyramid charges) is the most directly analogous precedent in the same category.")
    add_paragraph(doc,
        "Decision (Q1 2026): rebuild the public website from scratch on a compliance-first basis, before any new market campaign is approved.",
        bold=True)

    # 2. Problem Statement
    add_heading(doc, "2. Problem Statement", 1)
    add_paragraph(doc, "We need a public Tycoon Club website that:")
    for b in [
        "Can withstand a legal-and-compliance review without redaction of marketing claims.",
        "Does not rely on, promise, or imply guaranteed / passive / unlimited income.",
        "Removes the multi-tier referral compensation structure entirely.",
        "Gates qualified product information behind an application / eligibility step.",
        "Provides a single, deployable brand experience that the stakeholder team can choose from a set of clearly differentiated visual options.",
    ]:
        add_numbered(doc, b)
    add_paragraph(doc, "The brief from leadership:")
    add_quote(doc, "“Don’t tell people what they can earn. Tell people what we are, what we are not, and who can apply. Let counsel say yes before we ship a single growth campaign.”")

    # 3. Goals & Non-Goals
    add_heading(doc, "3. Goals & Non-Goals", 1)
    add_heading(doc, "3.1 Goals (in scope)", 2)
    for b in [
        "Compliance-first messaging — every page is built around what the product is and isn’t, with risk and eligibility surfaced before any product detail.",
        "Information-request form gate — users cannot self-enrol; they request an information packet, subject to acknowledgements and a jurisdictional check.",
        "Restricted-jurisdiction handling — the form refuses to submit from countries on the configured restricted list.",
        "Four parallel visual themes — same compliance backbone, four distinct design languages, so stakeholders can choose without re-doing copy.",
        "Theme picker — a single landing page that lets stakeholders preview and compare all four versions side-by-side.",
        "Static-deployable — no backend or build pipeline; deployable to any static host (Vercel, Netlify, GitHub Pages, S3+CloudFront).",
        "Mobile-responsive and accessibility-aware — works on phones, respects prefers-reduced-motion, semantic HTML.",
    ]:
        add_bullet(doc, b)
    add_heading(doc, "3.2 Non-goals (deferred)", 2)
    for b in [
        "Real membership onboarding or payment flow. The form is an information request, not an enrolment.",
        "A live contest engine, leaderboard, wallet connection, or NFT marketplace.",
        "KYC vendor integration. The form will store submissions client-side until a counsel-approved backend is wired up.",
        "Multi-language support. English-only for v1.",
        "A member portal or logged-in experience.",
        "Analytics / tracking pixels. To be added post-counsel-review with explicit cookie consent.",
    ]:
        add_bullet(doc, b)

    # 4. Target Users
    add_heading(doc, "4. Target Users", 1)
    add_table(doc,
        header=["Audience", "Use case"],
        rows=[
            ["Prospective member", "Lands from a search/referral, evaluates what Tycoon Club is, submits the information request."],
            ["Stakeholder / brand reviewer", "Compares the four themes to choose the brand direction."],
            ["Legal counsel", "Reviews the public-facing copy and the form acknowledgements before launch sign-off."],
            ["Compliance officer", "Validates the restricted-jurisdictions list and the risk-notice content."],
            ["Existing referrer", "Optionally includes their referral code in the form (one-time, capped — see §7)."],
        ],
        col_widths=[5.5, 11.0],
    )

    # 5. Functional Requirements
    add_heading(doc, "5. Functional Requirements", 1)

    add_heading(doc, "5.1 Pages (per version)", 2)
    add_table(doc,
        header=["Path", "Purpose"],
        rows=[
            ["/index.html", "Landing — compliant positioning, “what we are / are not”, member experience, referral framing, public notice."],
            ["/access.html", "Information-request form with acknowledgements and jurisdiction check. Post-submit reveals plain-language product summary."],
            ["/team.html", "Operating team roster (currently “Coming Soon” placeholders pending verifiable-identity policy)."],
            ["/compliance.html", "10-section risk notice, eligibility, restricted jurisdictions, terms summary, privacy."],
        ],
        col_widths=[4.0, 12.5],
    )

    add_heading(doc, "5.2 Information-request form (FR-01)", 2)
    add_paragraph(doc, "The form on /access.html is the load-bearing piece. Requirements:")
    for b in [
        "Required fields: full name, email, country of residence.",
        "Optional fields: free-text note. (A referrer field was present in earlier drafts and has been removed — the public site no longer surfaces a referral programme at all; see §7.3.)",
        "Five mandatory acknowledgements (submit button disabled until all five are checked):",
    ]:
        add_bullet(doc, b)
    for ack in [
        "18+ or local age of majority.",
        "Understands Tycoon Club is not an investment, securities offering, fund or yield product.",
        "Understands participation involves risk including loss of entry fees.",
        "Understands membership is subject to KYC and jurisdictional review.",
        "Accepts the Compliance Notice.",
    ]:
        p = doc.add_paragraph(style="List Number 2")
        p.add_run(ack).font.size = Pt(11)
    for b in [
        "Restricted-jurisdiction gate: if the selected country is on the restricted list (US, CN, KP, IR, SY, CU in v1), the form refuses submission and surfaces a “not available in your jurisdiction” notice with a deep-link to the Restricted Jurisdictions section of the Compliance Notice.",
        "Success state: the form swaps to a plain-English description of “what Tycoon Club really is” and “what it is not”, with three next-step bullets explaining the post-form workflow (information packet, KYC invitation, membership offer with full terms).",
        "No payment, no enrolment, no contest entry is offered through the form. This is documented in copy directly above the submit button.",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "5.3 Theme picker (FR-02)", 2)
    add_paragraph(doc, "A single page at /tycoon-club-themes/ that:")
    for b in [
        "Renders four preview tiles, each styled in its own theme’s actual palette and fonts so the visual choice is immediate.",
        "Each tile carries: a swatch row of the theme’s palette, the theme name, a 1-sentence positioning blurb, an “Enter →” button that opens the full site in a new tab, and a “Preview Inline” toggle that lazy-loads the site in an iframe inside the card.",
        "A comparison table at the bottom summarising what’s identical (compliance content, form behaviour) and what differs (fonts, palette, layout, feel).",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "5.4 Root landing (FR-03)", 2)
    add_paragraph(doc,
        "A small redirect page at / that immediately routes to /tycoon-club-themes/ via three layers (meta-refresh, JS location.replace, visible link fallback) so any client gets through.")

    # 6. Non-Functional Requirements
    add_heading(doc, "6. Non-Functional Requirements", 1)
    add_table(doc,
        header=["Category", "Requirement"],
        rows=[
            ["Compliance copy", "Every page must avoid the words guaranteed, passive, unlimited, lifetime, risk-free in any income context. The Compliance Notice explicitly flags these as red-flag language."],
            ["Mobile", "All breakpoints from 320px upward must be usable. Mobile nav menu, single-column form, responsive grid collapses."],
            ["Accessibility", "Semantic HTML, sufficient colour contrast for body copy, prefers-reduced-motion disables decorative animations (V02 gold dust, V04 hex orb float). 44×44px minimum tap targets for the mobile nav."],
            ["Performance", "No build step, no JS framework, no images larger than necessary. Each page loads in under ~50KB of CSS+JS. Google Fonts is the only external dependency."],
            ["Privacy", "No analytics, tracking pixels or third-party cookies in v1. Form submissions stay client-side until a counsel-reviewed backend lands."],
            ["No dependency lock-in", "Pure HTML/CSS/JS; portable to any static host."],
        ],
        col_widths=[4.5, 12.0],
    )

    # 7. Compliance backbone
    add_heading(doc, "7. Compliance Backbone (identical across all four versions)", 1)
    add_paragraph(doc, "This section captures what must not change regardless of the chosen visual theme.")

    add_heading(doc, "7.1 Positioning", 2)
    add_quote(doc,
        "Tycoon Club is a private members community for skill-based competitive gaming and digital entertainment. "
        "Membership is by application, subject to identity verification and jurisdictional review. We are not an "
        "investment scheme and do not offer guaranteed, “passive” or “unlimited” income.")
    add_paragraph(doc, "Every landing-page hero leads with that frame.")

    add_heading(doc, "7.2 “What we are / what we are not” dual block", 2)
    add_paragraph(doc, "Visible on every landing page above the fold (after the hero). Two columns:")
    add_paragraph(doc, "Tycoon Club is:", bold=True)
    for b in [
        "A members community for skill-based gaming and digital entertainment.",
        "An access platform for tournaments and leagues with published, transparent rules.",
        "Contests with prize pools funded from defined entry fees and sponsorships.",
        "Application-based with KYC, eligibility checks and jurisdictional review.",
    ]:
        add_bullet(doc, b)
    add_paragraph(doc, "Tycoon Club is not:", bold=True)
    for b in [
        "An investment, fund, securities offering, derivative or yield product.",
        "A source of guaranteed, “passive” or “unlimited” income.",
        "A public marketplace — access is restricted to approved members only.",
        "Available to residents of restricted jurisdictions or persons under 18.",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "7.3 Referral programme — public-site silence", 2)
    add_paragraph(doc,
        "The old site’s tiered commission structure was framed earlier in this project as a “one-time, capped community credit”. On reflection that framing still constituted a public claim about a referral compensation programme — and any public claim invites consistency-checks against the actual operational model. The public website is therefore silent on referral compensation in both directions: it does not advertise a referral programme, and it does not deny one. The “Community Growth Programme” landing-page section has been removed and replaced with a neutral “Member Privileges” block (curated events, private community, members-only content) modelled on standard private-club marketing. The optional “referrer” field has been removed from the access form, and the MLM-specific wording has been dropped from the form acknowledgements. Whatever the operational referral model is, it is governed by the member terms and operated outside the public website.")

    add_heading(doc, "7.4 Risk Notice / Compliance Ledger", 2)
    add_paragraph(doc, "Nine numbered sections on /compliance.html, deep-linked from every footer:")
    for i, t in enumerate([
        ("Not an investment", "explicit denial of securities, contracts, funds, yield products."),
        ("No guaranteed returns", "flags “guaranteed/passive/unlimited/lifetime/risk-free” as red-flag language."),
        ("Risk Notice", "entry fees at risk, digital asset volatility, past performance disclaimer."),
        ("Eligibility & KYC", "identity verification, sanctions/PEP screening, jurisdictional review, 18+."),
        ("Restricted Jurisdictions", "US, CN, KP, IR, SY, CU plus any comprehensive-sanctions country. List may be updated without notice."),
        ("Terms Summary", "full ToS governs contest rules, anti-cheat, dispute resolution, voiding."),
        ("Privacy", "data minimisation, withdrawal rights via compliance@tycoonclub.example."),
        ("Marketing & Communications", "no influencer income claims; unauthorised promotions to be reported."),
        ("Acknowledgement", "using the site constitutes acceptance."),
    ], start=1):
        add_numbered(doc, f"{t[0]} — {t[1]}")

    add_heading(doc, "7.5 Acknowledgements & jurisdictional block", 2)
    add_paragraph(doc, "The five-checkbox gate and the restricted-jurisdiction block in §5.2 are mandatory, present on every version, and identical in copy.")

    # 8. Solution overview
    add_heading(doc, "8. Solution overview — four themes, one backbone", 1)
    add_paragraph(doc,
        "We shipped four visually distinct versions, each carrying the same content from §7. Stakeholders can compare them via the theme picker and choose one (or compose a hybrid) before launch.")

    add_table(doc,
        header=["", "V01 CJC-aligned", "V02 Editorial", "V03 Fintech", "V04 Web3"],
        rows=[
            ["Inspiration", "CJC Race website (sibling site)", "“The Tycoon Brief” newsletter", "Premium crypto-exchange affiliate hub", "Smart-contract / Web3 affiliate sites"],
            ["Display font", "Orbitron", "Bodoni Moda (italic)", "Inter (bold)", "Space Grotesk"],
            ["Body font", "Rajdhani", "Inter Tight", "Inter", "Inter"],
            ["Background", "hsl(220 20% 6%)", "#0A0E14", "#0A0B0F", "#07081A"],
            ["Accent", "Bright orange-gold hsl(36 90% 55%)", "Antique gold #C9A961 + crimson #B33A3A", "Lime green #C5F84F + blue #3E80FF", "Cyan #00E0FF → violet #7B5BFF → pink #FF5BB8"],
            ["Signature layout", "Glass-card with backdrop blur, gold-gradient buttons", "Hairline cells, roman-numeral tags, animated gold-dust", "Bento grid, big stat numerals, pill buttons", "Hex icon chips, animated SVG hex orbs, neon glow"],
            ["Feel", "Gaming-forward, premium", "Editorial / luxury broadsheet", "Premium fintech / exchange", "Crypto-native / Web3"],
            ["Folder", "tycoon-club/", "tycoon-club-v2/", "tycoon-club-v3/", "tycoon-club-v4/"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5, 3.5],
    )

    add_heading(doc, "8.1 Vocabulary variants per theme", 2)
    add_paragraph(doc,
        "Some labels were shifted to match each theme’s voice. These are stylistic and carry no compliance weight — they can be reverted to a single canonical set before launch:")
    add_table(doc,
        header=["Concept", "V01 / V03 / V04", "V02"],
        rows=[
            ["Form CTA", "“Request Access”", "“Begin Your Petition”"],
            ["Team page title", "“Team” / “Meet the operating team”", "“The Operating Office”"],
            ["Compliance page title", "“Compliance Notice”", "“Compliance Ledger”"],
        ],
        col_widths=[5.0, 6.0, 5.5],
    )

    add_heading(doc, "8.2 Atmospheric / decorative elements", 2)
    add_table(doc,
        header=["Element", "Where", "Purpose", "Reduced-motion"],
        rows=[
            ["Glass-card backdrop blur", "V01 nav, V01 cards", "Premium-gaming feel", "n/a (static)"],
            ["Gold-dust radial-gradient particles, animated drift", "V02 body", "Editorial / heirloom feel", "Disabled"],
            ["Conic-gradient orb, floating stat cards", "V03 hero", "Fintech polish", "n/a (static)"],
            ["SVG hex orbs, float animation", "V04 hero", "Web3 / smart-contract vibe", "Disabled"],
            ["Dot-grid backdrop", "V04 body", "Crypto-native texture", "n/a (static)"],
        ],
        col_widths=[5.0, 3.5, 4.5, 3.5],
    )

    # 9. IA
    add_heading(doc, "9. Information architecture", 1)
    ia_paragraph = doc.add_paragraph()
    ia_run = ia_paragraph.add_run(
        "/\n"
        "├── /tycoon-club-themes/      Theme picker (preview, comparison table, deep-links)\n"
        "├── /tycoon-club/             V01\n"
        "│   ├── /access.html\n"
        "│   ├── /team.html\n"
        "│   └── /compliance.html\n"
        "├── /tycoon-club-v2/          V02 (same four pages, different skin)\n"
        "├── /tycoon-club-v3/          V03\n"
        "└── /tycoon-club-v4/          V04"
    )
    ia_run.font.name = "Consolas"
    ia_run.font.size = Pt(10)
    add_paragraph(doc,
        "Each version is self-contained: its own assets/styles.css and assets/script.js. Picking one for launch is a delete-three-folders operation.")

    # 10. Tech stack & deployment
    add_heading(doc, "10. Tech stack & deployment", 1)
    add_table(doc,
        header=["Layer", "Choice", "Rationale"],
        rows=[
            ["Markup", "Plain HTML5", "No framework lock-in, fastest TTFB, easiest for counsel to redline."],
            ["Styling", "Hand-rolled CSS (one file per version)", "Each theme is a self-contained stylesheet; no shared design tokens to avoid accidental cross-pollination."],
            ["Scripts", "Shared script.js (~60 LOC) per version", "Mobile-nav toggle, form acknowledgement-gating, restricted-jurisdiction block, localStorage submission persistence."],
            ["Fonts", "Google Fonts (Orbitron, Rajdhani, Bodoni Moda, Inter Tight, Inter, Space Grotesk, JetBrains Mono, IBM Plex Mono, Space Mono)", "Preconnect hints in every <head>."],
            ["Build", "None", "Static files; deployable with python -m http.server."],
            ["Hosting", "Vercel (auto-deploy on push to main)", "Configured via vercel.json with trailingSlash: true so directory index pages keep relative asset paths."],
            ["Repo", "GitHub GODKILLR/CJCtycoonclubwebsite", "Single main branch, conventional commits."],
        ],
        col_widths=[3.0, 5.5, 8.0],
    )

    add_heading(doc, "10.1 Local preview", 2)
    code = doc.add_paragraph()
    code_run = code.add_run("git clone https://github.com/GODKILLR/CJCtycoonclubwebsite\ncd CJCtycoonclubwebsite\npython -m http.server 5175\n# open http://localhost:5175/")
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(10)

    add_heading(doc, "10.2 Vercel configuration", 2)
    code2 = doc.add_paragraph()
    code2_run = code2.add_run("{ \"trailingSlash\": true }")
    code2_run.font.name = "Consolas"
    code2_run.font.size = Pt(10)
    add_paragraph(doc,
        "trailingSlash: true is load-bearing — without it, /tycoon-club-v3 serves index.html but the browser resolves relative assets/styles.css against /, breaking every directory index. With it on, Vercel redirects to /tycoon-club-v3/ before serving and paths resolve correctly.")

    # 11. Acceptance criteria
    add_heading(doc, "11. Acceptance criteria", 1)
    add_paragraph(doc, "A version is ready for launch when:")
    for n in [
        "Counsel sign-off on the Compliance Notice, the form acknowledgements, and the marketing & communications clause.",
        "Restricted-jurisdictions list is the legally-final list, not the v1 illustrative set.",
        "Operating entity disclosure replaces placeholder copy (name, registration number, registered address).",
        "compliance@tycoonclub.example is replaced with the real compliance mailbox and is monitored.",
        "Form submission backend is wired up — POST to a counsel-reviewed endpoint or vendor (Pipedream / HubSpot / Formspree / custom). The five-checkbox client-side gate stays as documented consent.",
        "Manual QA of mobile breakpoints (320px, 375px, 768px, 1280px) on the chosen theme.",
        "Accessibility pass: colour contrast (WCAG AA for body copy), keyboard navigation through the form, prefers-reduced-motion verification.",
    ]:
        add_numbered(doc, n)

    # 12. Open questions
    add_heading(doc, "12. Open questions / pre-launch checklist", 1)
    add_table(doc,
        header=["#", "Question", "Owner", "Status"],
        rows=[
            ["1", "Final operating entity name, registration number, registered address.", "Founders + counsel", "Open"],
            ["2", "Final restricted-jurisdictions list and sanctions-screening posture.", "Counsel + compliance", "Open"],
            ["3", "Whether to reintroduce any public reference to a referral programme. The public site is now silent on referral compensation (see §7.3). Any reintroduction — positive or negative — must be reviewed against the actual operational model before going live.", "Founders + counsel", "Open"],
            ["4", "Privacy / data-controller entity and full Privacy Notice.", "Counsel", "Open"],
            ["5", "Whether any contest formats trigger gambling-licence requirements in target markets.", "Counsel", "Open"],
            ["6", "Backend handover — which vendor receives form submissions and what is the data-retention policy?", "Engineering + compliance", "Open"],
            ["7", "Verifiable-identity policy for the Team page — when do we publish names, photos and LinkedIn links?", "Founders", "Open"],
            ["8", "Final theme selection.", "Brand + founders", "Pending stakeholder review"],
        ],
        col_widths=[1.0, 8.5, 4.0, 3.0],
    )

    # 13. Future work
    add_heading(doc, "13. Future work (post-v1)", 1)
    for b in [
        "Real backend for form intake + handoff to KYC vendor (e.g. Sumsub, Onfido, Persona).",
        "Member portal with logged-in experience, contest entry, leaderboard.",
        "Verified team profiles with on-chain identity attestations.",
        "Multi-language (priority: Arabic, Spanish, Portuguese, Mandarin if reopened).",
        "Analytics with explicit cookie consent (priority: Plausible or similar privacy-respecting tool).",
        "Press / media kit and a dedicated newsroom (the existing newsletter format can be ported).",
        "Compliance change log on /compliance.html so amendments are versioned.",
    ]:
        add_bullet(doc, b)

    # 14. Decision log
    add_heading(doc, "14. Decision log", 1)
    add_table(doc,
        header=["Date", "Decision", "Rationale"],
        rows=[
            ["2026-05", "Rebuild from scratch rather than redact the existing site",
                "The old structure (tiered referrals, achievement system) was load-bearing — fixing copy alone wouldn’t address the underlying MLM-shaped product."],
            ["2026-05", "Ship four parallel themes instead of one",
                "Stakeholders had not agreed a visual direction, and the prior site’s casino aesthetic was a contributing factor in regulatory risk. Parallel themes let leadership choose without re-doing content."],
            ["2026-05", "Pure static site, no framework",
                "Counsel can redline HTML/CSS directly; no build step, no runtime dependencies, lowest possible launch risk."],
            ["2026-05", "Form submissions to localStorage only in v1",
                "We did not want to wire a real intake endpoint before counsel had reviewed the form acknowledgements. Wiring happens after sign-off."],
            ["2026-05", "Restricted-jurisdiction list (US/CN/KP/IR/SY/CU) is illustrative",
                "The final list is counsel’s call; we surfaced a defensible starting point so the UX could be built end-to-end."],
            ["2026-05", "Drop the third-party brand name from the V03 theme label (renamed to “Fintech”)",
                "Naming a theme after a third-party brand is needlessly confusing and could imply affiliation."],
            ["2026-05", "Go silent on the referral programme on the public website",
                "The previous “one-time, capped community credit” framing was still a public claim about a referral compensation model. Any public claim invites consistency-checks against the actual operational model. The public site is now silent in both directions: it does not advertise a referral programme, and it does not deny one. The “Community Growth Programme” section is replaced with a neutral “Member Privileges” block; the referrer field is removed from the access form; the MLM-specific wording is removed from the form acknowledgements and the Compliance Notice section about referrals is deleted. Whatever the operational referral model is, it is governed by the member terms and operated outside the public site."],
        ],
        col_widths=[2.0, 5.5, 9.0],
    )

    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size:,} bytes)")
